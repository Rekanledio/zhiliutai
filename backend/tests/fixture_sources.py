from io import BytesIO

from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen.canvas import Canvas


def build_pdf_fixture() -> bytes:
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    output = BytesIO()
    canvas = Canvas(output, pagesize=letter)
    canvas.setTitle("合成 PDF 来源")
    canvas.setFont("STSong-Light", 14)
    canvas.drawString(72, 720, "知流台合成 PDF")
    canvas.setFont("STSong-Light", 11)
    canvas.drawString(72, 690, "第一页内容：本地优先知识库。")
    canvas.showPage()
    canvas.setFont("STSong-Light", 14)
    canvas.drawString(72, 720, "第二页")
    canvas.setFont("STSong-Light", 11)
    canvas.drawString(72, 690, "第二页内容：页码定位可追踪。")
    canvas.save()
    return output.getvalue()


def build_docx_fixture() -> bytes:
    document = Document()
    document.core_properties.title = "合成 DOCX 来源"
    document.add_heading("DOCX 知识指南", level=1)
    document.add_paragraph("第一段说明：文档标题层级需要保留。")
    document.add_heading("审核流程", level=2)
    document.add_paragraph("第二段说明：审核后再发布到 Obsidian。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "字段"
    table.rows[0].cells[1].text = "值"
    row = table.add_row().cells
    row[0].text = "来源"
    row[1].text = "合成 fixture"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def build_html_fixture() -> bytes:
    return """<!doctype html>
<html lang="zh-CN">
  <head>
    <title>静态网页合成指南</title>
    <script>throw new Error("must be ignored");</script>
  </head>
  <body>
    <nav>导航不应进入正文</nav>
    <main>
      <h1>静态网页指南</h1>
      <p>网页正文第一段，来源 URL 和标题层级需要保留。</p>
      <h2>安全边界</h2>
      <p>网页只允许无需登录的静态 HTML 来源。</p>
    </main>
    <footer>页脚不应进入正文</footer>
  </body>
</html>
""".encode("utf-8")
